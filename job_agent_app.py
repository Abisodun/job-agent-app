
import streamlit as st
import pandas as pd
import requests
from PyPDF2 import PdfReader
import docx
from docx import Document
import os
import openai

st.set_page_config(layout="wide")

RESUME_FILE = "resume.txt"
LOG_FILE = "job_log.csv"

client = openai.OpenAI()

def extract_text(file, file_type):
    if file_type == "txt":
        return file.read().decode("utf-8")
    elif file_type == "pdf":
        reader = PdfReader(file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_type == "docx":
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def score_resume_against_job(resume, job_description):
    resume_words = set(resume.lower().split())
    job_words = set(job_description.lower().split())
    match = resume_words.intersection(job_words)
    score = int((len(match) / len(job_words)) * 100) if job_words else 0
    return score

def score_to_stars(score):
    if score >= 90:
        return "⭐⭐⭐⭐⭐"
    elif score >= 75:
        return "⭐⭐⭐⭐"
    elif score >= 50:
        return "⭐⭐⭐"
    elif score >= 25:
        return "⭐⭐"
    else:
        return "⭐"

def fetch_jobs(query, country, city, employment_type, is_remote, min_salary, experience):
    url = "https://jsearch.p.rapidapi.com/search"
    headers = {
        "X-RapidAPI-Key": "2bc0998f49mshb8e8cd2a7c87e7dp18960fjsnf228ecfe1ab9",
        "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
    }
    full_query = f"{query or 'project manager'} in {city}, {country}"
    valid_types = ["Full-time", "Part-time", "Contract", "Temporary", "Internship"]

    params = {
        "query": full_query,
        "employment_types": employment_type if employment_type in valid_types else None,
        "remote_jobs_only": "true" if is_remote else "false",
        "min_salary": min_salary,
        "experience_level": experience.lower() if experience != "Any" else None,
        "page": "1",
        "num_pages": "2"
    }

    response = requests.get(url, headers=headers, params=params)
    if response.status_code != 200:
        st.error(f"API request failed with status code {response.status_code}")
        return []

    data = response.json()
    jobs = []
    for job in data.get("data", []):
        jobs.append({
            "Job Title": job.get("job_title", "N/A"),
            "Company": job.get("employer_name", "N/A"),
            "Location": job.get("job_city", "N/A") + ", " + job.get("job_country", "N/A"),
            "Link": job.get("job_apply_link", "N/A"),
            "Description": job.get("job_description", ""),
            "Cover Letter Path": "",
            "Status": "Applied"
        })
    return jobs

def generate_rewritten_resume(api_key, resume, job_description):
    client.api_key = api_key
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": f"Rewrite the following resume to better fit this job description. Use clear headings and bullet points for experience and skills.\n\nJob Description:\n{job_description}\n\nResume:\n{resume[:2500]}"}
            ],
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {e}"

def save_resume_as_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    filename = "Improved_Resume.docx"
    doc.save(filename)
    return filename

st.title("Smart Job Search + CV Rewrite Tool (v1.x compatible)")

if "api_key" not in st.session_state:
    api_key = st.text_input("Enter your OpenAI API Key", type="password")
    if api_key:
        st.session_state.api_key = api_key
        st.success("API key saved for session.")
else:
    st.success("✅ API Key is active.")

query = st.text_input("Job Title / Keywords", "project manager")
country = st.text_input("Country", "Canada")
city = st.text_input("City", "Toronto")
employment_type = st.selectbox("Employment Type", ["Any", "Full-time", "Part-time", "Contract", "Temporary", "Internship"])
experience = st.selectbox("Experience Level", ["Any", "Entry", "Mid", "Senior"])
is_remote = st.checkbox("Remote Only?")
min_salary = st.slider("Minimum Salary ($)", 0, 200000, 0, 5000)

uploaded_resume = st.file_uploader("Upload your Resume (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])
if uploaded_resume:
    file_type = uploaded_resume.name.split(".")[-1]
    resume_text = extract_text(uploaded_resume, file_type)
elif os.path.exists(RESUME_FILE):
    resume_text = open(RESUME_FILE, "r", encoding="utf-8").read()
else:
    st.warning("Please upload a resume or create a 'resume.txt' file.")
    st.stop()

if "jobs" not in st.session_state and st.button("Find & Rank Jobs"):
    with st.spinner("Ranking jobs by match score..."):
        st.session_state.jobs = fetch_jobs(query, country, city, employment_type, is_remote, min_salary, experience)
        for job in st.session_state.jobs:
            job["Match Score"] = score_resume_against_job(resume_text, job["Description"])
        st.session_state.jobs = sorted(st.session_state.jobs, key=lambda x: x["Match Score"], reverse=True)

if "jobs" in st.session_state:
    for i, job in enumerate(st.session_state.jobs[:10]):
        stars = score_to_stars(job["Match Score"])
        with st.expander(f"[{i+1}] {stars} {job['Job Title']} at {job['Company']} ({job['Match Score']}%)"):
            st.markdown(f"**Location:** {job['Location']}")
            st.markdown(f"[Apply Here]({job['Link']})")

            if st.button(f"Rewrite Resume for This Job", key=f"rewrite_{i}"):
                with st.spinner("Generating improved resume..."):
                    result = generate_rewritten_resume(st.session_state.api_key, resume_text, job["Description"])
                    if not result.startswith("Error"):
                        filename = save_resume_as_docx(result)
                        st.session_state.rewritten_resume_path = filename
                        st.session_state.rewritten_resume_text = result
                    else:
                        st.error(result)

            if st.session_state.get("rewritten_resume_text"):
                st.subheader("📄 Preview Rewritten Resume:")
                st.text_area("Preview", st.session_state.rewritten_resume_text, height=300)

            if st.session_state.get("rewritten_resume_path"):
                with open(st.session_state.rewritten_resume_path, "rb") as f:
                    st.download_button(
                        "Download Rewritten Resume (.docx)",
                        f,
                        file_name=st.session_state.rewritten_resume_path,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
